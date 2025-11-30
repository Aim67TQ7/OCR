"""
Text Agent - Enterprise Document Text Extraction
Strategic AI Advisory Implementation

Purpose: Autonomous text extraction agent for Claude Code integration.
Handles PDF, DOCX, TXT, and other text-based documents with enterprise-grade
error handling, metadata preservation, and performance optimization.
"""

import io
import logging
from typing import Dict, Any, Optional, List, Union
from pathlib import Path
import asyncio
from dataclasses import dataclass, asdict

# Core parsing dependencies
import PyPDF2
import pdfplumber
from docx import Document
import pandas as pd

@dataclass
class TextExtractionResult:
    """Standardized result container for text extraction operations."""
    text: str
    metadata: Dict[str, Any]
    extraction_method: str
    confidence_score: float
    processing_time_ms: int
    error_details: Optional[str] = None

class TextExtractionAgent:
    """
    Enterprise text extraction agent with autonomous decision-making capabilities.
    
    Strategic Value:
    - Autonomous method selection based on document characteristics
    - Performance-optimized extraction with fallback strategies
    - Enterprise metadata preservation for audit trails
    - Error isolation preventing pipeline failures
    """
    
    def __init__(self, 
                 optimization_mode: str = "balanced",
                 enable_fallbacks: bool = True,
                 max_text_length: Optional[int] = None):
        """
        Initialize text extraction agent.
        
        Args:
            optimization_mode: "speed", "quality", or "balanced"
            enable_fallbacks: Enable fallback extraction methods
            max_text_length: Maximum text length to prevent memory issues
        """
        self.optimization_mode = optimization_mode
        self.enable_fallbacks = enable_fallbacks
        self.max_text_length = max_text_length
        self.logger = logging.getLogger(__name__)
        
        # Performance tracking for strategic insights
        self.extraction_stats = {
            "total_extractions": 0,
            "method_usage": {},
            "avg_processing_time": 0,
            "success_rate": 0
        }

    async def extract_text(self, 
                          file_bytes: bytes, 
                          filename: Optional[str] = None,
                          mime_type: Optional[str] = None) -> TextExtractionResult:
        """
        Primary extraction interface - autonomously selects optimal method.
        
        Strategic Value: Single entry point with intelligent routing reduces
        integration complexity for enterprise applications.
        """
        start_time = asyncio.get_event_loop().time()
        
        try:
            # Autonomous method selection
            extraction_method = self._select_extraction_method(file_bytes, filename, mime_type)
            
            # Execute extraction with selected method
            result = await self._execute_extraction(file_bytes, extraction_method, filename)
            
            # Performance tracking
            processing_time = int((asyncio.get_event_loop().time() - start_time) * 1000)
            result.processing_time_ms = processing_time
            
            self._update_stats(extraction_method, processing_time, success=True)
            
            return result
            
        except Exception as e:
            self.logger.error(f"Text extraction failed: {str(e)}")
            processing_time = int((asyncio.get_event_loop().time() - start_time) * 1000)
            
            self._update_stats("error", processing_time, success=False)
            
            return TextExtractionResult(
                text="",
                metadata={"error": str(e)},
                extraction_method="failed",
                confidence_score=0.0,
                processing_time_ms=processing_time,
                error_details=str(e)
            )

    def _select_extraction_method(self, 
                                 file_bytes: bytes, 
                                 filename: Optional[str] = None,
                                 mime_type: Optional[str] = None) -> str:
        """
        Autonomous method selection based on document characteristics.
        
        Strategic Decision Tree:
        - PDF: pdfplumber for quality, PyPDF2 for speed
        - DOCX: python-docx native extraction
        - TXT: Direct UTF-8 decoding with fallbacks
        """
        
        # MIME type detection if not provided
        if not mime_type:
            mime_type = self._detect_mime_type(file_bytes, filename)
        
        # Strategic method selection
        if mime_type == "application/pdf":
            return "pdfplumber" if self.optimization_mode == "quality" else "pypdf2"
        elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            return "docx"
        elif mime_type.startswith("text/"):
            return "text_direct"
        else:
            # Fallback to binary analysis
            return self._analyze_binary_content(file_bytes)

    async def _execute_extraction(self, 
                                 file_bytes: bytes, 
                                 method: str, 
                                 filename: Optional[str] = None) -> TextExtractionResult:
        """Execute the selected extraction method with enterprise error handling."""
        
        extraction_methods = {
            "pdfplumber": self._extract_pdf_pdfplumber,
            "pypdf2": self._extract_pdf_pypdf2,
            "docx": self._extract_docx,
            "text_direct": self._extract_text_direct
        }
        
        primary_method = extraction_methods.get(method)
        if not primary_method:
            raise ValueError(f"Unsupported extraction method: {method}")
        
        try:
            result = await primary_method(file_bytes, filename)
            result.extraction_method = method
            return result
            
        except Exception as e:
            if self.enable_fallbacks:
                self.logger.warning(f"Primary method {method} failed, attempting fallback")
                return await self._attempt_fallback_extraction(file_bytes, method, filename)
            else:
                raise e

    async def _extract_pdf_pdfplumber(self, 
                                     file_bytes: bytes, 
                                     filename: Optional[str] = None) -> TextExtractionResult:
        """High-quality PDF text extraction using pdfplumber."""
        
        text_content = []
        metadata = {}
        
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            metadata = {
                "page_count": len(pdf.pages),
                "pdf_info": pdf.metadata or {},
                "extraction_method": "pdfplumber"
            }
            
            for page_num, page in enumerate(pdf.pages):
                try:
                    page_text = page.extract_text() or ""
                    text_content.append(page_text)
                except Exception as e:
                    self.logger.warning(f"Failed to extract page {page_num}: {str(e)}")
                    text_content.append("")
        
        full_text = "\n\n".join(text_content)
        
        # Apply length constraints
        if self.max_text_length and len(full_text) > self.max_text_length:
            full_text = full_text[:self.max_text_length]
            metadata["truncated"] = True
        
        return TextExtractionResult(
            text=full_text,
            metadata=metadata,
            extraction_method="pdfplumber",
            confidence_score=0.95,  # High confidence for pdfplumber
            processing_time_ms=0  # Will be set by caller
        )

    async def _extract_pdf_pypdf2(self, 
                                 file_bytes: bytes, 
                                 filename: Optional[str] = None) -> TextExtractionResult:
        """Fast PDF text extraction using PyPDF2."""
        
        text_content = []
        metadata = {}
        
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
        
        metadata = {
            "page_count": len(pdf_reader.pages),
            "pdf_info": pdf_reader.metadata or {},
            "extraction_method": "pypdf2"
        }
        
        for page_num, page in enumerate(pdf_reader.pages):
            try:
                page_text = page.extract_text()
                text_content.append(page_text)
            except Exception as e:
                self.logger.warning(f"Failed to extract page {page_num}: {str(e)}")
                text_content.append("")
        
        full_text = "\n\n".join(text_content)
        
        if self.max_text_length and len(full_text) > self.max_text_length:
            full_text = full_text[:self.max_text_length]
            metadata["truncated"] = True
        
        return TextExtractionResult(
            text=full_text,
            metadata=metadata,
            extraction_method="pypdf2",
            confidence_score=0.85,  # Lower confidence than pdfplumber
            processing_time_ms=0
        )

    async def _extract_docx(self, 
                           file_bytes: bytes, 
                           filename: Optional[str] = None) -> TextExtractionResult:
        """Extract text from DOCX documents."""
        
        doc = Document(io.BytesIO(file_bytes))
        
        text_content = []
        for paragraph in doc.paragraphs:
            text_content.append(paragraph.text)
        
        full_text = "\n".join(text_content)
        
        metadata = {
            "paragraph_count": len(doc.paragraphs),
            "extraction_method": "docx"
        }
        
        if self.max_text_length and len(full_text) > self.max_text_length:
            full_text = full_text[:self.max_text_length]
            metadata["truncated"] = True
        
        return TextExtractionResult(
            text=full_text,
            metadata=metadata,
            extraction_method="docx",
            confidence_score=0.98,  # Very high confidence for native format
            processing_time_ms=0
        )

    async def _extract_text_direct(self, 
                                  file_bytes: bytes, 
                                  filename: Optional[str] = None) -> TextExtractionResult:
        """Direct text extraction with encoding detection."""
        
        encodings = ['utf-8', 'latin-1', 'cp1252', 'ascii']
        text = ""
        used_encoding = None
        
        for encoding in encodings:
            try:
                text = file_bytes.decode(encoding)
                used_encoding = encoding
                break
            except UnicodeDecodeError:
                continue
        
        if not text:
            raise ValueError("Unable to decode text with any supported encoding")
        
        metadata = {
            "encoding": used_encoding,
            "file_size": len(file_bytes),
            "extraction_method": "text_direct"
        }
        
        if self.max_text_length and len(text) > self.max_text_length:
            text = text[:self.max_text_length]
            metadata["truncated"] = True
        
        return TextExtractionResult(
            text=text,
            metadata=metadata,
            extraction_method="text_direct",
            confidence_score=1.0,  # Perfect confidence for direct text
            processing_time_ms=0
        )

    async def _attempt_fallback_extraction(self, 
                                          file_bytes: bytes, 
                                          failed_method: str, 
                                          filename: Optional[str] = None) -> TextExtractionResult:
        """Fallback extraction strategies for enterprise resilience."""
        
        fallback_strategies = {
            "pdfplumber": ["pypdf2", "text_direct"],
            "pypdf2": ["pdfplumber", "text_direct"],
            "docx": ["text_direct"],
            "text_direct": []  # No fallback for direct text
        }
        
        fallbacks = fallback_strategies.get(failed_method, [])
        
        for fallback_method in fallbacks:
            try:
                result = await self._execute_extraction(file_bytes, fallback_method, filename)
                result.metadata["fallback_used"] = True
                result.metadata["original_method"] = failed_method
                result.confidence_score *= 0.8  # Reduce confidence for fallback
                return result
            except Exception as e:
                self.logger.warning(f"Fallback method {fallback_method} also failed: {str(e)}")
                continue
        
        raise ValueError(f"All extraction methods failed for {failed_method}")

    def _detect_mime_type(self, file_bytes: bytes, filename: Optional[str] = None) -> str:
        """Simple MIME type detection based on file signatures and extensions."""
        
        # Check file signatures (magic numbers)
        if file_bytes.startswith(b'%PDF'):
            return "application/pdf"
        elif file_bytes.startswith(b'PK\x03\x04') and filename and filename.endswith('.docx'):
            return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        
        # Fallback to filename extension
        if filename:
            suffix = Path(filename).suffix.lower()
            mime_map = {
                '.pdf': 'application/pdf',
                '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                '.txt': 'text/plain',
                '.md': 'text/markdown'
            }
            return mime_map.get(suffix, 'application/octet-stream')
        
        return 'application/octet-stream'

    def _analyze_binary_content(self, file_bytes: bytes) -> str:
        """Analyze binary content to determine best extraction approach."""
        
        # Simple heuristics for binary analysis
        if b'%PDF' in file_bytes[:1024]:
            return "pdfplumber"
        elif file_bytes.startswith(b'PK'):  # ZIP-like format (could be DOCX)
            return "docx"
        else:
            return "text_direct"

    def _update_stats(self, method: str, processing_time: int, success: bool):
        """Update performance statistics for strategic insights."""
        
        self.extraction_stats["total_extractions"] += 1
        
        if method not in self.extraction_stats["method_usage"]:
            self.extraction_stats["method_usage"][method] = {"count": 0, "avg_time": 0}
        
        method_stats = self.extraction_stats["method_usage"][method]
        method_stats["count"] += 1
        method_stats["avg_time"] = (method_stats["avg_time"] + processing_time) / 2
        
        if success:
            total_successful = sum(1 for m in self.extraction_stats["method_usage"].values() if m["count"] > 0)
            self.extraction_stats["success_rate"] = total_successful / self.extraction_stats["total_extractions"]

    def get_performance_insights(self) -> Dict[str, Any]:
        """
        Strategic performance insights for C-suite reporting.
        
        Returns operational metrics that translate to business value:
        - Processing efficiency by document type
        - Success rates and reliability metrics
        - Resource utilization patterns
        """
        
        return {
            "operational_summary": {
                "total_documents_processed": self.extraction_stats["total_extractions"],
                "overall_success_rate": f"{self.extraction_stats['success_rate']:.1%}",
                "optimization_mode": self.optimization_mode
            },
            "method_performance": self.extraction_stats["method_usage"],
            "strategic_recommendations": self._generate_strategic_recommendations()
        }

    def _generate_strategic_recommendations(self) -> List[str]:
        """Generate actionable strategic recommendations based on usage patterns."""
        
        recommendations = []
        method_usage = self.extraction_stats.get("method_usage", {})
        
        if not method_usage:
            return ["Insufficient data for recommendations"]
        
        # Most used method
        most_used = max(method_usage.items(), key=lambda x: x[1]["count"])
        recommendations.append(f"Primary extraction method: {most_used[0]} (used {most_used[1]['count']} times)")
        
        # Performance optimization suggestions
        if self.optimization_mode == "balanced":
            recommendations.append("Consider 'speed' mode for high-volume processing or 'quality' mode for critical documents")
        
        # Fallback usage insights
        if any("fallback_used" in str(method_usage) for method_usage in method_usage.values()):
            recommendations.append("Fallback methods in use - consider document quality assessment")
        
        return recommendations

# Strategic Usage Example for C-suite Integration
async def demonstrate_enterprise_usage():
    """
    Example demonstrating enterprise-grade text extraction
    for strategic AI advisory implementations.
    """
    
    # Initialize agent with enterprise configuration
    agent = TextExtractionAgent(
        optimization_mode="balanced",
        enable_fallbacks=True,
        max_text_length=1_000_000  # 1MB text limit
    )
    
    # Example usage (would be called by Claude Code)
    # result = await agent.extract_text(file_bytes, filename="contract.pdf")
    
    # Strategic insights for executive reporting
    insights = agent.get_performance_insights()
    
    return {
        "agent_ready": True,
        "enterprise_features": [
            "Autonomous method selection",
            "Fallback strategies",
            "Performance tracking",
            "Strategic insights generation"
        ],
        "integration_points": [
            "Claude Code agent system",
            "Enterprise document pipelines",
            "Executive reporting dashboards"
        ]
    }
